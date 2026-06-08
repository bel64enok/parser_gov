from __future__ import annotations

import html
import io
import json
import mimetypes
import os
import re
import shutil
import socket
import tempfile
import threading
import time
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime, timedelta
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception:  # pragma: no cover - optional runtime dependency
    Document = None

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional runtime dependency
    load_workbook = None


ROOT = Path(__file__).resolve().parent
STATIC_DIR = ROOT / "static"
UPLOAD_DIR = ROOT / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

ZAKUPKI_BASE = "https://zakupki.gov.ru"


DOMAIN_TRIGGERS: dict[str, list[str]] = {
    "МОБ": [
        "мобильная связь",
        "сотовая связь",
        "сим-карт",
        "sim",
        "смс",
        "sms",
        "lte",
        "4g",
        "5g",
        "радиотелефон",
        "подвижной связи",
        "подвижная связь",
        "голосовая связь",
        "мобильный интернет",
    ],
    "ФИКС": [
        "доступ в интернет",
        "фиксированная связь",
        "канал связи",
        "vpn",
        "волс",
        "ip-телефони",
        "телематическ",
    ],
    "SI": ["системная интеграция", "оборудование", "монтаж", "пусконалад"],
    "М2М": ["m2m", "iot", "телеметр", "датчик", "мониторинг транспорта"],
    "BigData": ["big data", "большие данные", "аналитика данных", "скоринг"],
    "Cloud": ["облач", "iaas", "saas", "виртуальн", "цод", "хостинг"],
    "SBVAS": ["vas", "контент", "массовые вызовы", "короткий номер", "ivr"],
}

PRODUCT_CATALOG = {
    "Корпоративная мобильная связь": ["мобильная связь", "сотовая связь", "радиотелефон", "sim", "сим-карт"],
    "Мобильный интернет": ["мобильный интернет", "lte", "4g", "5g"],
    "SMS и уведомления": ["sms", "смс", "коротк", "рассыл"],
    "Фиксированный интернет": ["доступ в интернет", "канал связи", "волс"],
    "Облачные сервисы": ["облач", "виртуальн", "iaas", "saas"],
}

SAMPLE_TENDERS = [
    {
        "number": "0373200123426000311",
        "title": "Оказание услуг подвижной радиотелефонной связи и мобильного доступа в интернет",
        "customer": "ГБУ города Москвы",
        "price": 8_450_000,
        "method": "Электронный аукцион",
        "law": "44-ФЗ",
        "publish_date": "02.06.2026",
        "deadline": "11.06.2026",
        "appeared_at": "05.06.2026 09:20",
        "contract_status": "Исполнение",
        "okpd2": "61.20.11.000",
        "status": "Исполнение",
        "url": "https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString=0373200123426000311",
    },
    {
        "number": "0173100003226000042",
        "title": "Оказание услуг фиксированного доступа к сети Интернет и каналов связи",
        "customer": "Федеральное казенное учреждение",
        "price": 16_900_000,
        "method": "Запрос котировок",
        "law": "44-ФЗ",
        "publish_date": "03.06.2026",
        "deadline": "09.06.2026",
        "appeared_at": "04.06.2026 15:45",
        "contract_status": "Исполнение",
        "okpd2": "61.10.30.190",
        "status": "Исполнение",
        "url": "https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString=0173100003226000042",
    },
    {
        "number": "32515229187",
        "title": "Поставка облачной платформы для обработки данных и виртуальных машин",
        "customer": "АО Региональный оператор",
        "price": 42_000_000,
        "method": "Конкурс в электронной форме",
        "law": "223-ФЗ",
        "publish_date": "01.06.2026",
        "deadline": "19.06.2026",
        "appeared_at": "01.06.2026 11:10",
        "contract_status": "Завершен",
        "okpd2": "63.11.19.000",
        "status": "Завершен",
        "url": "https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString=32515229187",
    },
]


def sample_tenders() -> list[dict[str, Any]]:
    tenders = [dict(item) for item in SAMPLE_TENDERS]
    templates = [
        ("Оказание услуг мобильной связи для территориальных подразделений", "МОБ", "Электронный аукцион", "61.20.11.000", 6_900_000),
        ("Предоставление каналов связи и фиксированного доступа в интернет", "ФИКС", "Электронный аукцион", "61.10.30.190", 12_400_000),
        ("Услуги M2M мониторинга транспорта и SIM-карт для телеметрии", "М2М", "Запрос котировок", "61.20.30.000", 4_700_000),
        ("Услуги SMS-информирования и коротких сообщений", "SBVAS", "Электронный аукцион", "61.20.42.000", 3_300_000),
        ("Корпоративная мобильная связь и мобильный интернет", "МОБ", "Электронный аукцион", "61.20.11.000", 9_850_000),
        ("Фиксированные VPN-каналы связи между объектами заказчика", "ФИКС", "Запрос котировок", "61.10.30.190", 5_600_000),
        ("Передача данных IoT-устройств и телеметрии", "М2М", "Электронный аукцион", "61.20.30.000", 7_100_000),
        ("Подключение мобильных номеров и SIM-карт для сотрудников", "МОБ", "Запрос котировок", "61.20.11.000", 2_900_000),
        ("Организация защищенных каналов связи", "ФИКС", "Электронный аукцион", "61.10.30.190", 18_200_000),
        ("Услуги голосовой связи и мобильного интернета", "МОБ", "Электронный аукцион", "61.20.11.000", 14_500_000),
        ("SMS-рассылки для уведомления граждан", "SBVAS", "Запрос котировок", "61.20.42.000", 1_700_000),
        ("Сотовая связь для оперативных служб", "МОБ", "Электронный аукцион", "61.20.11.000", 22_000_000),
        ("Интернет-каналы для филиальной сети", "ФИКС", "Электронный аукцион", "61.10.30.190", 31_000_000),
        ("Сервис мониторинга датчиков с передачей данных", "М2М", "Запрос котировок", "61.20.30.000", 3_950_000),
        ("Мобильная связь и SMS-пакеты", "МОБ", "Электронный аукцион", "61.20.11.000", 8_750_000),
    ]
    for index, (title, domain, method, okpd2, price) in enumerate(templates, start=4):
        day = min(27, 5 + index)
        tenders.append(
            {
                "number": f"037320012342600{index:04d}",
                "title": title,
                "customer": f"Заказчик демо-пула {index}",
                "price": price,
                "method": method,
                "law": "44-ФЗ",
                "publish_date": f"{max(1, day - 7):02d}.06.2026",
                "deadline": f"{day:02d}.06.2026",
                "appeared_at": f"{max(1, 5 - (index % 3)):02d}.06.2026 {9 + (index % 8):02d}:15",
                "contract_status": "Исполнение",
                "okpd2": okpd2,
                "status": "Исполнение",
                "url": f"https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString=037320012342600{index:04d}",
                "demo_domain": domain,
            }
        )
    return tenders


@dataclass
class ChecklistItem:
    name: str
    status: str
    result: str
    comment: str


@dataclass
class EvidenceItem:
    category: str
    label: str
    term: str
    snippet: str


def lower_text(value: str) -> str:
    return (value or "").lower().replace("ё", "е")


def parse_money(value: str | int | float | None) -> int:
    if isinstance(value, (int, float)):
        return int(value)
    if not value:
        return 0
    clean = re.sub(r"[^\d,\.]", "", value)
    if not clean:
        return 0
    clean = clean.replace(",", ".")
    if clean.count(".") > 1:
        clean = clean.replace(".", "")
    return int(float(clean))


def parse_date(value: str | None) -> datetime | None:
    if not value:
        return None
    for fmt in ("%d.%m.%Y", "%d.%m.%Y %H:%M", "%Y-%m-%d"):
        try:
            return datetime.strptime(value.strip()[:16], fmt)
        except ValueError:
            pass
    return None


def days_between_exclusive(start: str | None, end: str | None, business: bool = False) -> int | None:
    start_dt = parse_date(start)
    end_dt = parse_date(end)
    if not start_dt or not end_dt or end_dt <= start_dt:
        return None
    current = start_dt.date() + timedelta(days=1)
    last = end_dt.date() - timedelta(days=1)
    count = 0
    while current <= last:
        if not business or current.weekday() < 5:
            count += 1
        current += timedelta(days=1)
    return count


def detect_domains(text: str) -> list[str]:
    haystack = lower_text(text)
    domains = []
    for domain, triggers in DOMAIN_TRIGGERS.items():
        if any(lower_text(trigger) in haystack for trigger in triggers):
            domains.append(domain)
    return domains


def detect_products(text: str) -> list[str]:
    haystack = lower_text(text)
    return [
        product
        for product, triggers in PRODUCT_CATALOG.items()
        if any(lower_text(trigger) in haystack for trigger in triggers)
    ]


def document_text_for_tender(tender: dict[str, Any]) -> str:
    """Prototype substitute for downloaded tender documents."""
    title = tender.get("title") or "Без названия"
    method = tender.get("method") or "Способ закупки не распознан"
    okpd2 = tender.get("okpd2") or "ОКПД2 не указан"
    price = tender.get("price") or 0
    publish_date = tender.get("publish_date") or "Дата размещения не распознана"
    deadline = tender.get("deadline") or "Дата окончания не распознана"
    customer = tender.get("customer") or "Заказчик не распознан"
    status = tender.get("status") or "Статус не распознан"
    return (
        f"Извещение о закупке № {tender.get('number', '')}.\n"
        f"Наименование объекта закупки: {title}.\n"
        f"Заказчик: {customer}.\n"
        f"Способ определения поставщика: {method}.\n"
        f"Начальная максимальная цена контракта: {price} руб.\n"
        f"Код ОКПД2: {okpd2}.\n"
        f"Дата размещения извещения: {publish_date}.\n"
        f"Дата и время окончания подачи заявок: {deadline}.\n"
        f"Статус контракта: {tender.get('contract_status') or status}.\n"
        "Документ используется прототипом как распознанный текст карточки и прикрепленной документации."
    )


def evidence_snippet(text: str, term: str, radius: int = 150) -> str:
    normalized_text = lower_text(text)
    normalized_term = lower_text(term)
    index = normalized_text.find(normalized_term)
    if index < 0:
        return ""
    start = max(0, index - radius)
    end = min(len(text), index + len(term) + radius)
    snippet = text[start:end].strip()
    return re.sub(r"\s+", " ", snippet)


def collect_evidence(tender: dict[str, Any], document_text: str) -> list[dict[str, str]]:
    evidence: list[EvidenceItem] = []
    seen: set[tuple[str, str, str]] = set()

    def add(category: str, label: str, term: str, source_text: str = document_text) -> None:
        snippet = evidence_snippet(source_text, term)
        key = (category, label, term)
        if snippet and key not in seen:
            evidence.append(EvidenceItem(category, label, term, snippet))
            seen.add(key)

    for domain, triggers in DOMAIN_TRIGGERS.items():
        for trigger in triggers:
            add("Домен", domain, trigger)

    method = str(tender.get("method", ""))
    for term in ["аукцион", "котиров", "конкурс", "способ определения поставщика"]:
        add("Способ закупки", method or "Способ закупки", term)

    for term in ["ОКПД2", str(tender.get("okpd2", "")), "61.", "начальная максимальная цена", "цена контракта"]:
        if term:
            add("Критерии", "ОКПД2 и сумма", term)

    for term in ["дата размещения", "окончания подачи заявок", str(tender.get("publish_date", "")), str(tender.get("deadline", ""))]:
        if term:
            add("Критерии", "Сроки подачи", term)

    return [asdict(item) for item in evidence]


def documents_for_tender(tender: dict[str, Any]) -> list[dict[str, Any]]:
    text = document_text_for_tender(tender)
    source_url = tender.get("url") or build_search_url(str(tender.get("number", "")))
    return [
        {
            "id": "notice",
            "name": "Извещение и распознанный текст карточки",
            "type": "Извещение ЕИС",
            "text": text,
            "source_url": source_url,
            "highlights": collect_evidence(tender, text),
        },
        {
            "id": "eis-documents",
            "name": "Прикрепленные документы на сайте госзакупок",
            "type": "Документы ЕИС",
            "text": (
                "Откройте источник ЕИС, чтобы посмотреть и скачать оригинальные файлы, "
                "прикрепленные к карточке тендера на сайте госзакупок.\n\n"
                f"Номер закупки: {tender.get('number', '')}.\n"
                f"Наименование: {tender.get('title', '')}.\n"
                f"Источник: {source_url}"
            ),
            "source_url": source_url,
            "highlights": [],
        }
    ]


def analyze_tender(tender: dict[str, Any], document_text: str = "") -> dict[str, Any]:
    combined = " ".join(
        str(tender.get(key, ""))
        for key in ("title", "customer", "method", "okpd2", "status", "law")
    )
    combined = f"{combined} {document_text}"
    domains = detect_domains(combined)
    products = detect_products(combined)
    price = parse_money(tender.get("price"))
    method_text = lower_text(str(tender.get("method", "")))
    okpd2 = str(tender.get("okpd2", ""))
    checklist: list[ChecklistItem] = []

    if domains:
        status = "pass" if "МОБ" in domains else "info"
        result = "Найден домен МОБ, анализ запущен" if "МОБ" in domains else "МОБ не найден"
        checklist.append(
            ChecklistItem(
                "Определение домена",
                status,
                result,
                "Домены: " + ", ".join(domains),
            )
        )
    else:
        checklist.append(
            ChecklistItem(
                "Определение домена",
                "warning",
                "Домен не определен",
                "Добавьте ключевые слова или загрузите документацию.",
            )
        )

    method_ok = True
    method_comment = []
    if okpd2 and not okpd2.startswith("61"):
        method_comment.append("ОКПД2 не относится к телеком-услугам 61.*")
    if "котиров" in method_text and price > 10_000_000:
        method_ok = False
        method_comment.append("Запрос котировок выше 10 млн руб. требует проверки")
    if "конкурс" in method_text:
        method_ok = False
        method_comment.append("Конкурс по заданным критериям помечается как неверный способ")
    if "аукцион" in method_text:
        method_comment.append("Аукцион допустим при любой сумме")
    checklist.append(
        ChecklistItem(
            "Определение способа закупки",
            "pass" if method_ok else "risk",
            "Способ закупки выглядит корректно" if method_ok else "Способ закупки требует внимания",
            "; ".join(method_comment) or "Недостаточно данных по способу закупки.",
        )
    )

    required_days = None
    actual_days = None
    business = False
    if "аукцион" in method_text:
        required_days = 15 if price > 300_000_000 else 7
    elif "котиров" in method_text:
        required_days = 4
        business = True
    elif "конкурс" in method_text:
        required_days = 15
    if required_days:
        actual_days = days_between_exclusive(tender.get("publish_date"), tender.get("deadline"), business)
    if required_days and actual_days is not None:
        deadline_ok = actual_days >= required_days
        unit = "рабочих" if business else "календарных"
        checklist.append(
            ChecklistItem(
                "Сроки подачи заявок",
                "pass" if deadline_ok else "risk",
                f"{actual_days} из требуемых {required_days} {unit} дней",
                "Срок считается без даты размещения и даты окончания подачи.",
            )
        )
    else:
        checklist.append(
            ChecklistItem(
                "Сроки подачи заявок",
                "warning",
                "Недостаточно дат для проверки",
                "Нужны дата размещения, дата окончания и способ закупки.",
            )
        )

    risk_count = sum(1 for item in checklist if item.status == "risk")
    warning_count = sum(1 for item in checklist if item.status == "warning")
    score = 42
    if "МОБ" in domains:
        score += 24
    if okpd2.startswith("61"):
        score += 14
    if products:
        score += 8
    score -= risk_count * 18
    score -= warning_count * 6
    if price and price <= 10_000_000:
        score += 4
    score = max(0, min(100, score))

    priority = "Высокий" if score >= 75 else "Средний" if score >= 50 else "Низкий"
    return {
        "domains": domains,
        "products": products,
        "checklist": [asdict(item) for item in checklist],
        "evidence": collect_evidence(tender, combined),
        "risk_count": risk_count,
        "warning_count": warning_count,
        "score": score,
        "priority": priority,
        "queue_priority": tender.get("deadline") or "Без срока",
    }


def build_search_url(query: str, page: int = 1, limit: int = 10, refresh_token: str | None = None) -> str:
    params = {
        "searchString": query,
        "morphology": "on",
        "search-filter": "Дате размещения",
        "pageNumber": str(page),
        "sortDirection": "false",
        "recordsPerPage": f"_{limit}",
        "showLotsInfoHidden": "false",
        "sortBy": "UPDATE_DATE",
        "fz44": "on",
        "fz223": "on",
        "pc": "on",
        "currencyIdGeneral": "-1",
    }
    if refresh_token:
        params["_refresh"] = refresh_token
    return f"{ZAKUPKI_BASE}/epz/order/extendedsearch/results.html?{urllib.parse.urlencode(params)}"


def fetch_zakupki_html(query: str, limit: int, refresh_token: str | None = None) -> tuple[str, str]:
    url = build_search_url(query, limit=limit, refresh_token=refresh_token)
    request = urllib.request.Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 TenderParserPrototype/0.1",
            "Accept-Language": "ru-RU,ru;q=0.9,en;q=0.8",
            "Cache-Control": "no-cache",
            "Pragma": "no-cache",
        },
    )
    with urllib.request.urlopen(request, timeout=14) as response:
        charset = response.headers.get_content_charset() or "utf-8"
        return response.read().decode(charset, errors="ignore"), url


def strip_tags(value: str) -> str:
    value = re.sub(r"<script[\s\S]*?</script>", " ", value, flags=re.I)
    value = re.sub(r"<style[\s\S]*?</style>", " ", value, flags=re.I)
    value = re.sub(r"<[^>]+>", " ", value)
    value = html.unescape(value)
    return re.sub(r"\s+", " ", value).strip()


def find_near(text: str, patterns: list[str]) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, re.I | re.S)
        if match:
            return strip_tags(match.group(1))
    return ""


def parse_zakupki_results(page_html: str) -> list[dict[str, Any]]:
    chunks = re.split(r'class="[^"]*(?:search-registry-entry-block|registry-entry__form)[^"]*"', page_html)
    if len(chunks) <= 1:
        chunks = re.split(r"(?=№\s*\d{11,})", page_html)
    tenders = []
    for raw in chunks[1:]:
        chunk = raw[:20000]
        text = strip_tags(chunk)
        number = find_near(chunk, [r"№\s*</span>\s*<a[^>]*>\s*([^<]+)", r"№\s*([0-9]{8,})"])
        if not number:
            number_match = re.search(r"\b\d{11,19}\b", text)
            number = number_match.group(0) if number_match else ""
        if not number:
            continue
        title = find_near(
            chunk,
            [
                r'class="[^"]*registry-entry__body-value[^"]*"[^>]*>\s*<a[^>]*>([\s\S]*?)</a>',
                r'<a[^>]+href="[^"]*common-info[^"]*"[^>]*>([\s\S]*?)</a>',
            ],
        )
        if not title:
            title = re.sub(r"^.*?№\s*" + re.escape(number), "", text)[:220].strip()
        href_match = re.search(r'href="([^"]*common-info[^"]*)"', chunk)
        url = urllib.parse.urljoin(ZAKUPKI_BASE, href_match.group(1)) if href_match else build_search_url(number)
        price = find_near(
            chunk,
            [
                r"Начальная[^<]{0,80}</[^>]+>\s*<[^>]+>([\s\S]*?)</",
                r"Цена контракта[^<]{0,80}</[^>]+>\s*<[^>]+>([\s\S]*?)</",
            ],
        )
        publish_date = find_near(chunk, [r"Размещено\s*</[^>]+>\s*<[^>]+>([\d\.]+)"])
        deadline = find_near(chunk, [r"Окончание подачи[^<]*</[^>]+>\s*<[^>]+>([\d\.\s:]+)"])
        customer = find_near(chunk, [r"Заказчик[^<]*</[^>]+>\s*<[^>]+>([\s\S]*?)</"])
        method = find_near(chunk, [r"Способ определения[^<]*</[^>]+>\s*<[^>]+>([\s\S]*?)</"])
        tenders.append(
            {
                "number": number.strip(),
                "title": title or "Закупка без распознанного названия",
                "customer": customer or "Не распознано",
                "price": parse_money(price),
                "method": method or "Не распознано",
                "law": "44-ФЗ" if "44-ФЗ" in text else "223-ФЗ" if "223-ФЗ" in text else "ЕИС",
                "publish_date": publish_date,
                "deadline": deadline[:10] if deadline else "",
                "okpd2": "",
                "status": "Исполнение" if "Исполнение" in text else "Подача заявок" if "Подача заявок" in text else "",
                "contract_status": "Исполнение" if "Исполнение" in text else "Подача заявок" if "Подача заявок" in text else "",
                "appeared_at": publish_date,
                "url": url,
            }
        )
    return tenders[:20]


def search_tenders(query: str, limit: int = 10, refresh_token: str | None = None) -> dict[str, Any]:
    query = query.strip() or "мобильная связь"
    started = time.monotonic()
    parsed_at = datetime.now().isoformat(timespec="seconds")
    live_count = 0
    try:
        page_html, source_url = fetch_zakupki_html(query, limit, refresh_token=refresh_token)
        tenders = parse_zakupki_results(page_html)
        live_count = len(tenders)
        if tenders:
            source = "zakupki.gov.ru"
            error = ""
        else:
            tenders = sample_tenders()
            source = "fallback"
            error = "ЕИС ответила, но структура страницы не распознана. Показаны демо-записи."
    except (urllib.error.URLError, socket.timeout, TimeoutError, OSError) as exc:
        source_url = build_search_url(query, limit=limit, refresh_token=refresh_token)
        tenders = sample_tenders()
        source = "fallback"
        error = f"ЕИС недоступна или ответила слишком медленно: {exc.__class__.__name__}."

    enriched = []
    for tender in tenders[:limit]:
        if tender.get("contract_status") != "Исполнение":
            continue
        item = dict(tender)
        item["documents"] = documents_for_tender(item)
        item["analysis"] = analyze_tender(item, "\n".join(document["text"] for document in item["documents"]))
        enriched.append(item)
    if not enriched and source != "fallback":
        source = "fallback"
        error = "ЕИС ответила, но карточки со статусом «Исполнение» не распознаны. Показаны демо-записи."
        for tender in sample_tenders()[:limit]:
            if tender.get("contract_status") != "Исполнение":
                continue
            item = dict(tender)
            item["documents"] = documents_for_tender(item)
            item["analysis"] = analyze_tender(item, "\n".join(document["text"] for document in item["documents"]))
            enriched.append(item)
    enriched.sort(key=lambda item: (-item["analysis"]["score"], item.get("deadline") or "99.99.9999"))
    return {
        "items": enriched,
        "source": source,
        "source_url": source_url,
        "error": error,
        "query": query,
        "parsed_at": parsed_at,
        "parse_duration_ms": int((time.monotonic() - started) * 1000),
        "live_count": live_count,
        "refresh_token": refresh_token or "",
    }


def extract_docx(path: Path) -> str:
    if Document is None:
        return ""
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def extract_xlsx(path: Path) -> str:
    if load_workbook is None:
        return ""
    workbook = load_workbook(path, read_only=True, data_only=True)
    values = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            values.extend(str(cell) for cell in row if cell is not None)
    return "\n".join(values)


def extract_pdf_text_layer(path: Path) -> str:
    data = path.read_bytes()
    text = re.sub(rb"[^\x20-\x7E\xD0-\xFF]+", b" ", data[:2_000_000])
    return text.decode("utf-8", errors="ignore")


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            return extract_docx(path)
        if suffix in {".xlsx", ".xlsm"}:
            return extract_xlsx(path)
        if suffix == ".pdf":
            return extract_pdf_text_layer(path)
        if suffix in {".txt", ".csv"}:
            return path.read_text("utf-8", errors="ignore")
        if suffix == ".zip":
            parts = []
            with zipfile.ZipFile(path) as archive:
                with tempfile.TemporaryDirectory() as tmp:
                    tmpdir = Path(tmp)
                    archive.extractall(tmpdir)
                    for child in tmpdir.rglob("*"):
                        if child.is_file() and child.suffix.lower() in {".docx", ".xlsx", ".xlsm", ".pdf", ".txt", ".csv"}:
                            parts.append(extract_text_from_file(child))
            return "\n".join(parts)
    except Exception as exc:
        return f"Ошибка извлечения {path.name}: {exc}"
    return ""


def requirements_payload() -> dict[str, Any]:
    return {
        "functional": [
            "Загрузка комплектов документации в систему: zip, docx, xlsx, pdf, txt.",
            "Распознавание документов DOCX, Excel и PDF с текстовым слоем; OCR отмечен как следующий этап.",
            "Выделение ключевых формулировок, тегов доменов и продуктовой маркировки.",
            "Обучаемый словарь доменов и продуктов для добавления новых категорий.",
            "Web-интерфейс загрузки, поиска, просмотра карточек, ранжирования и критериев анализа.",
        ],
        "non_functional": [
            "Целевое SLA обработки одного тендера: до 8 часов.",
            "Очередь минимум 50 тендеров/час с приоритетом по дате окончания подачи заявок.",
            "Доступность web-интерфейса: целевой уровень 99% в месяц.",
            "Конфиденциальность данных и развертывание в инфраструктуре заказчика.",
            "Готовность к расширению критериев анализа и актуализации законодательства.",
        ],
        "checklist": [
            "Определение домена по словам-триггерам: МОБ, ФИКС, SI, М2М, BigData, Cloud, SBVAS.",
            "Проверка способа закупки: ОКПД2 61.*, сумма до 10 млн для запроса котировок, конкурс требует внимания.",
            "Проверка сроков подачи: аукцион 7/15 календарных дней, котировки 4 рабочих дня, конкурс 15 календарных дней.",
        ],
        "domains": DOMAIN_TRIGGERS,
    }


def _parse_multipart(rfile: io.RawIOBase, headers: object) -> dict:
    """Minimal multipart/form-data parser replacing the removed cgi.FieldStorage."""
    content_type: str = headers.get("Content-Type", "")
    content_length = int(headers.get("Content-Length", 0) or 0)
    boundary: bytes | None = None
    for chunk in content_type.split(";"):
        chunk = chunk.strip()
        if chunk.startswith("boundary="):
            boundary = chunk[9:].strip().encode()
            break
    if not boundary:
        return {}
    body = rfile.read(content_length)
    fields: dict = {}
    for raw_part in body.split(b"--" + boundary)[1:]:
        if raw_part.lstrip(b"\r\n").startswith(b"--"):
            break
        if b"\r\n\r\n" not in raw_part:
            continue
        raw_headers, content = raw_part.split(b"\r\n\r\n", 1)
        if content.endswith(b"\r\n"):
            content = content[:-2]
        name: str | None = None
        filename: str | None = None
        for line in raw_headers.decode("utf-8", errors="replace").splitlines():
            lower = line.lower()
            if "content-disposition" in lower:
                for item in line.split(";"):
                    item = item.strip()
                    if item.startswith("name="):
                        name = item[5:].strip('"')
                    elif item.startswith("filename="):
                        filename = item[9:].strip('"')
        if name:
            fields[name] = type("_Field", (), {"filename": filename, "file": io.BytesIO(content)})()
    return fields


class TenderHandler(BaseHTTPRequestHandler):
    server_version = "TenderParserPrototype/0.1"

    def log_message(self, fmt: str, *args: Any) -> None:
        print("[%s] %s" % (self.log_date_time_string(), fmt % args))

    def send_json(self, payload: Any, status: int = 200) -> None:
        data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        self.wfile.write(data)

    def do_GET(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/api/health":
            self.send_json({"status": "ok", "time": datetime.now().isoformat(timespec="seconds")})
            return
        if parsed.path == "/api/requirements":
            self.send_json(requirements_payload())
            return
        if parsed.path == "/api/tenders":
            query = urllib.parse.parse_qs(parsed.query)
            search = query.get("query", ["мобильная связь"])[0]
            limit = int(query.get("limit", ["10"])[0])
            refresh_token = query.get("refresh", [str(int(time.time() * 1000))])[0]
            self.send_json(search_tenders(search, max(1, min(limit, 100)), refresh_token=refresh_token))
            return
        self.serve_static(parsed.path)

    def do_POST(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path != "/api/upload":
            self.send_json({"error": "Not found"}, 404)
            return
        form = _parse_multipart(self.rfile, self.headers)
        upload = form.get("file")
        if upload is None or not getattr(upload, "filename", ""):
            self.send_json({"error": "Файл не передан"}, 400)
            return
        safe_name = re.sub(r"[^A-Za-zА-Яа-я0-9_.() -]+", "_", Path(upload.filename).name)
        target = UPLOAD_DIR / f"{int(time.time())}_{safe_name}"
        with target.open("wb") as file_obj:
            shutil.copyfileobj(upload.file, file_obj)
        text = extract_text_from_file(target)
        virtual_tender = {
            "number": "LOCAL-" + datetime.now().strftime("%Y%m%d%H%M%S"),
            "title": safe_name,
            "customer": "Загруженный комплект документации",
            "price": 0,
            "method": "",
            "law": "Локальный анализ",
            "publish_date": "",
            "deadline": "",
            "okpd2": "61" if "61" in text else "",
            "status": "В обработке",
            "url": "",
        }
        analysis = analyze_tender(virtual_tender, text)
        document = {
            "id": "upload-" + datetime.now().strftime("%Y%m%d%H%M%S"),
            "name": safe_name,
            "type": target.suffix.lower().lstrip(".") or "file",
            "text": text,
            "highlights": collect_evidence(virtual_tender, text),
        }
        self.send_json(
            {
                "file": safe_name,
                "stored_as": str(target),
                "characters": len(text),
                "analysis": analysis,
                "document": document,
                "preview": text[:1200],
            }
        )

    def serve_static(self, path: str) -> None:
        if path in {"", "/"}:
            path = "/index.html"
        target = (STATIC_DIR / path.lstrip("/")).resolve()
        if not str(target).startswith(str(STATIC_DIR.resolve())) or not target.exists() or not target.is_file():
            self.send_json({"error": "Not found"}, 404)
            return
        content = target.read_bytes()
        mime = mimetypes.guess_type(str(target))[0] or "application/octet-stream"
        self.send_response(200)
        self.send_header("Content-Type", mime + ("; charset=utf-8" if mime.startswith("text/") or mime == "application/javascript" else ""))
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)


def run(host: str = "127.0.0.1", port: int = 8000) -> None:
    httpd = ThreadingHTTPServer((host, port), TenderHandler)
    print(f"Tender Parser prototype: http://{host}:{port}")
    httpd.serve_forever()


if __name__ == "__main__":
    run(port=int(os.environ.get("PORT", "8000")))
