"""Извлечение текста из документов: docx / xlsx / pdf (текстовый слой) / txt / csv / zip.

Перенесено из app.py без изменения логики. Используется стадией 2 (analyzer) и /api/upload.
"""
from __future__ import annotations

import re
import tempfile
import zipfile
from pathlib import Path

try:
    from docx import Document
except Exception:  # pragma: no cover - optional runtime dependency
    Document = None

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional runtime dependency
    load_workbook = None


def extract_docx(path: Path) -> str:
    if Document is None:
        return ""
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def extract_xlsx(path: Path) -> str:
    if load_workbook is None:
        return ""
    workbook = load_workbook(path, read_only=True, data_only=True)
    values = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            values.extend(str(cell) for cell in row if cell is not None)
    return "\n".join(values)


def extract_pdf_text_layer(path: Path) -> str:
    data = path.read_bytes()
    text = re.sub(rb"[^\x20-\x7E\xD0-\xFF]+", b" ", data[:2_000_000])
    return text.decode("utf-8", errors="ignore")


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            return extract_docx(path)
        if suffix in {".xlsx", ".xlsm"}:
            return extract_xlsx(path)
        if suffix == ".pdf":
            return extract_pdf_text_layer(path)
        if suffix in {".txt", ".csv"}:
            return path.read_text("utf-8", errors="ignore")
        if suffix in {".html", ".htm"}:
            from scrape import strip_tags  # локальный импорт во избежание цикла

            return strip_tags(path.read_text("utf-8", errors="ignore"))
        if suffix == ".zip":
            parts = []
            with zipfile.ZipFile(path) as archive:
                with tempfile.TemporaryDirectory() as tmp:
                    tmpdir = Path(tmp)
                    archive.extractall(tmpdir)
                    for child in tmpdir.rglob("*"):
                        if child.is_file() and child.suffix.lower() in {".docx", ".xlsx", ".xlsm", ".pdf", ".txt", ".csv"}:
                            parts.append(extract_text_from_file(child))
            return "\n".join(parts)
    except Exception as exc:
        return f"Ошибка извлечения {path.name}: {exc}"
    return ""
